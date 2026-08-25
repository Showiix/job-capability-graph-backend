import asyncio
import re
from dataclasses import dataclass
from datetime import date
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from pypdf import PdfReader

from app.core.errors import APIError
from app.resumes.schemas import ResumeParseResponse

PDF_MEDIA_TYPES = {"application/pdf", "application/octet-stream"}
DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
DOCX_MEDIA_TYPES = {DOCX_MEDIA_TYPE, "application/octet-stream"}
MAX_RESUME_FILE_BYTES = 20 * 1024 * 1024
MAX_DOCX_UNCOMPRESSED_BYTES = 100 * 1024 * 1024
MAX_EXTRACTED_TEXT_CHARS = 100_000

HORIZONTAL_WHITESPACE = re.compile(r"[^\S\n]+")
PHONE_PATTERN = re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")
EMAIL_PATTERN = re.compile(
    r"[A-Z0-9._%+-]{1,64}@[A-Z0-9-]{1,63}(?:\.[A-Z0-9-]{1,63})+",
    re.IGNORECASE,
)
IDENTITY_PATTERN = re.compile(
    r"(?<!\d)\d{6}(?:18|19|20)\d{2}(?:0[1-9]|1[0-2])"
    r"(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx](?!\d)"
)
WECHAT_PATTERN = re.compile(
    r"(?P<label>(?:微信号?|WeChat)[ \t]*[:：]?[ \t]*)"
    r"(?P<value>[A-Z][A-Z0-9_-]{5,19})",
    re.IGNORECASE,
)

EDUCATION_RANK = {
    "unknown": 0,
    "other": 1,
    "high_school": 2,
    "associate": 3,
    "bachelor": 4,
    "master": 5,
    "doctor": 6,
}
EVIDENCE_RANK = {"mention": 0, "project": 1, "work": 2}


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    text: str
    method: str


@dataclass(frozen=True, slots=True)
class ValidatedParse:
    document_language: str
    summary: str | None
    educations: list[dict]
    experiences: list[dict]
    projects: list[dict]
    skills: list[dict]
    warnings: list[str]


def detect_resume_document(filename: str, media_type: str, content: bytes) -> str:
    extension = Path(filename).suffix.lower()
    declared_type = media_type.split(";", 1)[0].strip().lower()

    if extension == ".pdf":
        if declared_type not in PDF_MEDIA_TYPES or not content.startswith(b"%PDF-"):
            raise _unsupported_file_type()
        return "pdf"
    if extension == ".docx":
        if declared_type not in DOCX_MEDIA_TYPES:
            raise _unsupported_file_type()
        validate_docx_archive(content)
        return "docx"
    raise _unsupported_file_type()


def validate_docx_archive(content: bytes) -> None:
    try:
        with ZipFile(BytesIO(content)) as archive:
            infos = archive.infolist()
            names = {info.filename for info in infos}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValueError("required docx entries missing")
            if any(info.flag_bits & 0x1 for info in infos):
                raise ValueError("encrypted docx entry")
            if sum(info.file_size for info in infos) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("docx uncompressed size exceeded")
    except (BadZipFile, OSError, ValueError) as error:
        raise APIError(422, "RESUME_DOCUMENT_INVALID", "简历文档结构无效") from error


async def extract_resume_text(path: Path, document_type: str) -> ExtractedDocument:
    try:
        if document_type == "pdf":
            raw_text = await asyncio.to_thread(_extract_pdf_text, path)
            method = "pdf_text"
        elif document_type == "docx":
            raw_text = await asyncio.to_thread(_extract_docx_text, path)
            method = "docx"
        else:
            raise ValueError("unknown resume document type")
    except APIError:
        raise
    except Exception as error:
        raise APIError(422, "RESUME_DOCUMENT_INVALID", "简历文档结构无效") from error

    return ExtractedDocument(text=normalize_extracted_text(raw_text), method=method)


def normalize_extracted_text(value: str) -> str:
    normalized_newlines = value.replace("\r\n", "\n").replace("\r", "\n")
    lines: list[str] = []
    previous_blank = False
    for source_line in normalized_newlines.split("\n"):
        line = HORIZONTAL_WHITESPACE.sub(" ", source_line).strip()
        if line:
            lines.append(line)
            previous_blank = False
        elif lines and not previous_blank:
            lines.append("")
            previous_blank = True

    normalized = "\n".join(lines).strip()
    if not normalized:
        raise APIError(422, "RESUME_TEXT_EMPTY", "简历中未提取到可用文本")
    if len(normalized) > MAX_EXTRACTED_TEXT_CHARS:
        raise APIError(422, "RESUME_TEXT_TOO_LONG", "简历正文超过处理上限")
    return normalized


def redact_resume_text(value: str) -> str:
    redacted = PHONE_PATTERN.sub(_stars, value)
    redacted = EMAIL_PATTERN.sub(_stars, redacted)
    redacted = IDENTITY_PATTERN.sub(_stars, redacted)
    return WECHAT_PATTERN.sub(_redact_wechat, redacted)


def locate_evidence(text: str, quote: str) -> tuple[int, int] | None:
    if not quote:
        return None
    start = text.find(quote)
    if start < 0:
        return None
    return start, start + len(quote)


def validate_parse_evidence(
    payload: ResumeParseResponse,
    *,
    redacted_text: str,
) -> ValidatedParse:
    warnings: list[str] = []
    educations = _ground_items(
        payload.educations,
        category="EDUCATION",
        label_field="school_name",
        redacted_text=redacted_text,
        warnings=warnings,
    )
    experiences = _ground_items(
        payload.experiences,
        category="EXPERIENCE",
        label_field="company_name",
        redacted_text=redacted_text,
        warnings=warnings,
    )
    projects = _ground_items(
        payload.projects,
        category="PROJECT",
        label_field="project_name",
        redacted_text=redacted_text,
        warnings=warnings,
    )
    skills = _ground_items(
        payload.skills,
        category="SKILL",
        label_field="name",
        redacted_text=redacted_text,
        warnings=warnings,
    )
    if not any((educations, experiences, projects, skills)):
        raise APIError(422, "RESUME_EVIDENCE_EMPTY", "简历解析结果缺少可追溯证据")
    return ValidatedParse(
        document_language=payload.document_language,
        summary=payload.summary,
        educations=educations,
        experiences=experiences,
        projects=projects,
        skills=skills,
        warnings=warnings,
    )


def skill_rank(skill: dict) -> tuple[int, float, int]:
    return (
        EVIDENCE_RANK[skill["evidence_strength"]],
        float(skill["confidence"]),
        -int(skill["evidence_start"]),
    )


def derive_highest_education(educations: list[dict]) -> str | None:
    if not educations:
        return None
    levels = [
        value if value in EDUCATION_RANK else "unknown"
        for item in educations
        if (value := item.get("education_level")) is not None
    ]
    if not levels:
        return "unknown"
    return max(levels, key=EDUCATION_RANK.__getitem__)


def derive_total_experience_months(
    experiences: list[dict],
    *,
    current_month: date,
) -> tuple[int | None, list[str]]:
    intervals: list[tuple[int, int]] = []
    warnings: list[str] = []
    current_index = current_month.year * 12 + current_month.month - 1

    for item in experiences:
        start_month = item.get("start_month")
        end_month = item.get("end_month")
        is_current = item.get("is_current") is True
        if not start_month or (not is_current and not end_month):
            warnings.append("EXPERIENCE_DATE_INCOMPLETE")
            continue
        try:
            start_index = _month_index(start_month)
            end_index = current_index if is_current else _month_index(end_month)
        except (TypeError, ValueError):
            warnings.append("EXPERIENCE_DATE_INCOMPLETE")
            continue
        if end_index < start_index:
            warnings.append("EXPERIENCE_DATE_INCOMPLETE")
            continue
        intervals.append((start_index, end_index + 1))

    if not intervals:
        return None, warnings

    merged: list[list[int]] = []
    for start, end in sorted(intervals):
        if not merged or start > merged[-1][1]:
            merged.append([start, end])
        else:
            merged[-1][1] = max(merged[-1][1], end)
    return sum(end - start for start, end in merged), warnings


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(path: Path) -> str:
    content = path.read_bytes()
    validate_docx_archive(content)
    document = Document(BytesIO(content))
    values = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            values.extend(cell.text for cell in row.cells)
    return "\n".join(values)


def _month_index(value: str) -> int:
    year_text, month_text = value.split("-", 1)
    year = int(year_text)
    month = int(month_text)
    if len(year_text) != 4 or len(month_text) != 2 or not 1 <= month <= 12:
        raise ValueError("invalid month")
    return year * 12 + month - 1


def _ground_items(
    items: list,
    *,
    category: str,
    label_field: str,
    redacted_text: str,
    warnings: list[str],
) -> list[dict]:
    grounded = []
    for item in items:
        value = item.model_dump(mode="python")
        offsets = locate_evidence(redacted_text, value["evidence_quote"])
        if (
            offsets is None
            and category == "SKILL"
            and value[label_field] in value["evidence_quote"]
        ):
            offsets = locate_evidence(redacted_text, value[label_field])
            if offsets is not None:
                value["evidence_quote"] = value[label_field]
        if offsets is None:
            warnings.append(f"{category}_EVIDENCE_NOT_FOUND:{value[label_field]}")
            continue
        value["evidence_start"], value["evidence_end"] = offsets
        grounded.append(value)
    return grounded


def _stars(match: re.Match[str]) -> str:
    return "*" * len(match.group(0))


def _redact_wechat(match: re.Match[str]) -> str:
    return f"{match.group('label')}{'*' * len(match.group('value'))}"


def _unsupported_file_type() -> APIError:
    return APIError(415, "RESUME_FILE_TYPE_UNSUPPORTED", "仅支持 PDF 或 DOCX 简历")
