from datetime import date
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

import pytest
from docx import Document
from pypdf import PdfWriter

from app.core.errors import APIError
from app.resumes.parsing import (
    DOCX_MEDIA_TYPE,
    derive_highest_education,
    derive_total_experience_months,
    detect_resume_document,
    extract_resume_text,
    locate_evidence,
    normalize_extracted_text,
    redact_resume_text,
    validate_docx_archive,
    validate_parse_evidence,
)
from app.resumes.schemas import ResumeParseResponse


def make_docx_bytes() -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("Python 项目经验")
    document.save(stream)
    return stream.getvalue()


def test_pdf_requires_pdf_signature() -> None:
    assert detect_resume_document("resume.pdf", "application/pdf", b"%PDF-1.7") == "pdf"
    with pytest.raises(APIError) as error:
        detect_resume_document("resume.pdf", "application/pdf", b"not-pdf")
    assert error.value.code == "RESUME_FILE_TYPE_UNSUPPORTED"


def test_docx_requires_office_zip_entries() -> None:
    assert (
        detect_resume_document("resume.docx", DOCX_MEDIA_TYPE, make_docx_bytes())
        == "docx"
    )
    stream = BytesIO()
    with ZipFile(stream, "w", ZIP_DEFLATED) as archive:
        archive.writestr("word/other.xml", "<xml/>")

    with pytest.raises(APIError) as error:
        validate_docx_archive(stream.getvalue())
    assert error.value.code == "RESUME_DOCUMENT_INVALID"


def test_resume_images_require_matching_signature() -> None:
    assert (
        detect_resume_document("resume.jpg", "image/jpeg", b"\xff\xd8\xffdemo")
        == "image"
    )
    assert (
        detect_resume_document("resume.png", "image/png", b"\x89PNG\r\n\x1a\ndemo")
        == "image"
    )
    with pytest.raises(APIError):
        detect_resume_document("resume.png", "image/png", b"not-png")


def test_declared_unrelated_media_type_is_rejected() -> None:
    with pytest.raises(APIError) as error:
        detect_resume_document("resume.pdf", "image/png", b"%PDF-1.7")

    assert error.value.status_code == 415


def test_octet_stream_is_allowed_only_when_signature_matches() -> None:
    assert (
        detect_resume_document("resume.pdf", "application/octet-stream", b"%PDF-1.7")
        == "pdf"
    )


def test_corrupt_docx_zip_is_rejected() -> None:
    with pytest.raises(APIError) as error:
        validate_docx_archive(b"not-a-zip")

    assert error.value.code == "RESUME_DOCUMENT_INVALID"


def test_encrypted_docx_entry_is_rejected(monkeypatch) -> None:
    def encrypted_infos(_archive):
        content_types = ZipInfo("[Content_Types].xml")
        document = ZipInfo("word/document.xml")
        document.flag_bits = 0x1
        return [content_types, document]

    monkeypatch.setattr(ZipFile, "infolist", encrypted_infos)
    with pytest.raises(APIError) as error:
        validate_docx_archive(make_docx_bytes())

    assert error.value.code == "RESUME_DOCUMENT_INVALID"


def test_docx_uncompressed_size_is_bounded(monkeypatch) -> None:
    monkeypatch.setattr("app.resumes.parsing.MAX_DOCX_UNCOMPRESSED_BYTES", 10)

    with pytest.raises(APIError) as error:
        validate_docx_archive(make_docx_bytes())

    assert error.value.code == "RESUME_DOCUMENT_INVALID"


async def test_extracts_text_pdf_fixture() -> None:
    path = Path(__file__).parent / "fixtures" / "resume_text.pdf"

    result = await extract_resume_text(path, "pdf")

    assert result.method == "pdf_text"
    assert "Python" in result.text


async def test_extracts_docx_paragraphs_and_table(tmp_path) -> None:
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("示例大学")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "FastAPI"
    document.save(path)

    result = await extract_resume_text(path, "docx")

    assert result.method == "docx"
    assert result.text.splitlines() == ["示例大学", "Python", "FastAPI"]


def test_normalize_text_keeps_words_and_normalizes_whitespace() -> None:
    assert normalize_extracted_text("A\r\n\r\n  B\t C  ") == "A\n\nB C"


def test_empty_text_is_rejected() -> None:
    with pytest.raises(APIError) as error:
        normalize_extracted_text(" \n\t ")

    assert error.value.code == "RESUME_TEXT_EMPTY"


def test_text_length_boundary() -> None:
    assert len(normalize_extracted_text("汉" * 100_000)) == 100_000
    with pytest.raises(APIError) as error:
        normalize_extracted_text("汉" * 100_001)

    assert error.value.code == "RESUME_TEXT_TOO_LONG"


async def test_corrupt_documents_are_classified(tmp_path) -> None:
    pdf_path = tmp_path / "bad.pdf"
    pdf_path.write_bytes(b"%PDF-invalid")
    docx_path = tmp_path / "bad.docx"
    docx_path.write_bytes(b"not-a-docx")

    for path, document_type in ((pdf_path, "pdf"), (docx_path, "docx")):
        with pytest.raises(APIError) as error:
            await extract_resume_text(path, document_type)
        assert error.value.code == "RESUME_DOCUMENT_INVALID"


async def test_blank_pdf_is_rejected_as_empty_text(tmp_path) -> None:
    path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(path)

    with pytest.raises(APIError) as error:
        await extract_resume_text(path, "pdf")

    assert error.value.code == "RESUME_TEXT_EMPTY"


def test_redaction_is_length_preserving() -> None:
    original = (
        "手机：13800138000\n"
        "邮箱：demo@example.com\n"
        "身份证：110101199001011234\n"
        "微信号：demo_wechat-1\n"
        "Python 项目"
    )

    redacted = redact_resume_text(original)

    assert len(redacted) == len(original)
    assert redacted.count("\n") == original.count("\n")
    assert "13800138000" not in redacted
    assert "demo@example.com" not in redacted
    assert "110101199001011234" not in redacted
    assert "demo_wechat-1" not in redacted
    assert "Python 项目" in redacted


@pytest.mark.parametrize(
    ("text", "secret"),
    [
        ("电话 13800138000 完成", "13800138000"),
        ("邮箱 demo@example.com 完成", "demo@example.com"),
        ("身份证 110101199001011234 完成", "110101199001011234"),
        ("微信号：demo_wechat-1 完成", "demo_wechat-1"),
    ],
)
def test_each_pii_pattern_only_replaces_the_value(text, secret) -> None:
    assert redact_resume_text(text) == text.replace(secret, "*" * len(secret))


def test_redaction_does_not_replace_names_or_ordinary_tokens() -> None:
    text = "姓名：张三\nPython developer token"

    assert redact_resume_text(text) == text


def test_evidence_offsets_use_unicode_code_points() -> None:
    original = "甲乙\n使用 Python 开发项目"
    redacted = redact_resume_text(original)

    evidence = locate_evidence(redacted, "使用 Python 开发项目")

    assert evidence == (3, 17)
    assert original[evidence[0] : evidence[1]] == "使用 Python 开发项目"


def test_repeated_quote_uses_first_occurrence() -> None:
    assert locate_evidence("Python\n其他\nPython", "Python") == (0, 6)


def test_missing_quote_returns_none() -> None:
    assert locate_evidence("Python", "Java") is None


def test_highest_education_uses_backend_order() -> None:
    assert (
        derive_highest_education(
            [{"education_level": "bachelor"}, {"education_level": "master"}]
        )
        == "master"
    )
    assert derive_highest_education([]) is None


def test_experience_months_merge_overlapping_closed_intervals() -> None:
    experiences = [
        {"start_month": "2024-01", "end_month": "2024-03", "is_current": False},
        {"start_month": "2024-03", "end_month": "2024-05", "is_current": False},
    ]

    total, warnings = derive_total_experience_months(
        experiences, current_month=date(2026, 8, 1)
    )

    assert total == 5
    assert warnings == []


def test_ongoing_experience_uses_worker_utc_month() -> None:
    total, warnings = derive_total_experience_months(
        [{"start_month": "2026-06", "end_month": None, "is_current": True}],
        current_month=date(2026, 8, 1),
    )

    assert total == 3
    assert warnings == []


def test_incomplete_experience_is_kept_but_not_counted() -> None:
    total, warnings = derive_total_experience_months(
        [{"start_month": None, "end_month": None, "is_current": False}],
        current_month=date(2026, 8, 1),
    )

    assert total is None
    assert warnings == ["EXPERIENCE_DATE_INCOMPLETE"]


def test_invalid_evidence_items_are_dropped_with_warnings() -> None:
    payload = ResumeParseResponse.model_validate(
        {
            "schema_version": "resume_parse_v1",
            "document_language": "zh-CN",
            "summary": "示例",
            "educations": [],
            "experiences": [],
            "projects": [],
            "skills": [
                {
                    "name": "Python",
                    "proficiency": None,
                    "explicit_experience_months": None,
                    "evidence_strength": "project",
                    "evidence_quote": "使用 Python 开发项目",
                    "confidence": 0.9,
                },
                {
                    "name": "Java",
                    "proficiency": None,
                    "explicit_experience_months": None,
                    "evidence_strength": "mention",
                    "evidence_quote": "不存在的 Java 证据",
                    "confidence": 0.8,
                },
            ],
        }
    )

    validated = validate_parse_evidence(
        payload,
        redacted_text="使用 Python 开发项目",
    )

    assert [item["name"] for item in validated.skills] == ["Python"]
    assert validated.skills[0]["evidence_start"] == 0
    assert validated.skills[0]["evidence_end"] == 14
    assert validated.warnings == ["SKILL_EVIDENCE_NOT_FOUND:Java"]


def test_skill_name_is_safe_evidence_fallback() -> None:
    payload = ResumeParseResponse.model_validate(
        {
            "schema_version": "resume_parse_v1",
            "document_language": "zh-CN",
            "summary": None,
            "educations": [],
            "experiences": [],
            "projects": [],
            "skills": [
                {
                    "name": "FastAPI",
                    "proficiency": "intermediate",
                    "explicit_experience_months": None,
                    "evidence_strength": "project",
                    "evidence_quote": "使用 FastAPI 开发接口",
                    "confidence": 0.9,
                }
            ],
        }
    )

    validated = validate_parse_evidence(
        payload,
        redacted_text="使用 Python 和 FastAPI 开发接口",
    )

    assert validated.skills[0]["evidence_quote"] == "FastAPI"
    assert validated.warnings == []


@pytest.mark.parametrize("with_candidates", [False, True])
def test_evidence_gate_rejects_when_no_category_has_grounded_item(
    with_candidates,
) -> None:
    values = {
        "schema_version": "resume_parse_v1",
        "document_language": "zh-CN",
        "summary": "只有摘要不能通过",
        "educations": [],
        "experiences": [],
        "projects": [],
        "skills": [],
    }
    if with_candidates:
        values["educations"] = [
            {
                "school_name": "示例大学",
                "major": None,
                "education_level": "bachelor",
                "start_month": None,
                "end_month": None,
                "is_current": False,
                "evidence_quote": "不存在的学历",
                "confidence": 0.8,
            }
        ]
        values["experiences"] = [
            {
                "company_name": "示例公司",
                "job_title": None,
                "start_month": None,
                "end_month": None,
                "is_current": False,
                "responsibilities": [],
                "evidence_quote": "不存在的经历",
                "confidence": 0.8,
            }
        ]
        values["projects"] = [
            {
                "project_name": "示例项目",
                "role": None,
                "start_month": None,
                "end_month": None,
                "is_current": False,
                "description": None,
                "evidence_quote": "不存在的项目",
                "confidence": 0.8,
            }
        ]
        values["skills"] = [
            {
                "name": "Python",
                "proficiency": None,
                "explicit_experience_months": None,
                "evidence_strength": "mention",
                "evidence_quote": "不存在的技能",
                "confidence": 0.8,
            }
        ]
    payload = ResumeParseResponse.model_validate(values)

    with pytest.raises(APIError) as error:
        validate_parse_evidence(payload, redacted_text="无相关正文")

    assert error.value.code == "RESUME_EVIDENCE_EMPTY"
