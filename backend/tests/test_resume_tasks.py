from datetime import UTC, datetime
from hashlib import sha256
from types import SimpleNamespace
from uuid import uuid4

import pytest
import pytest_asyncio
from docx import Document
from pydantic import SecretStr
from sqlalchemy import func, select
from sqlalchemy.exc import SQLAlchemyError

from app.catalog.models import Capability, CapabilityAlias, Domain
from app.files.models import StoredFile
from app.infrastructure.file_storage import FileStorage
from app.processing.models import ProcessingError, ProcessingRun
from app.processing.service import retry_run
from app.resumes.llm import LLMParseResult, ResumeLLMError
from app.resumes.models import Resume, ResumeProfile, ResumeSkill
from app.resumes.parsing import DOCX_MEDIA_TYPE
from app.resumes.schemas import ResumeParseResponse
from app.resumes.service import map_resume_skills
from app.resumes.tasks import run_parse_resume


async def add_capability(
    db_session,
    *,
    domain_name: str,
    canonical_name: str,
    status: str = "active",
    alias: str | None = None,
    alias_status: str = "active",
) -> Capability:
    domain = Domain(
        id=uuid4(),
        code=f"domain-{uuid4().hex}",
        name=domain_name,
        status="active",
        sort_order=0,
    )
    capability = Capability(
        id=uuid4(),
        domain_id=domain.id,
        canonical_name=canonical_name,
        skill_type="technical",
        status=status,
        source_type="manual",
    )
    db_session.add(domain)
    await db_session.flush()
    db_session.add(capability)
    await db_session.flush()
    if alias is not None:
        db_session.add(
            CapabilityAlias(
                id=uuid4(),
                capability_id=capability.id,
                alias=alias,
                status=alias_status,
            )
        )
        await db_session.flush()
    return capability


def skill(
    name: str,
    *,
    strength: str = "mention",
    confidence: float = 0.8,
    start: int = 0,
) -> dict:
    return {
        "name": name,
        "proficiency": None,
        "explicit_experience_months": None,
        "evidence_strength": strength,
        "evidence_quote": name,
        "evidence_start": start,
        "evidence_end": start + len(name),
        "confidence": confidence,
    }


async def capability_count(db_session) -> int:
    value = await db_session.scalar(select(func.count()).select_from(Capability))
    return int(value or 0)


async def test_canonical_exact_maps_only_one_active_capability(db_session) -> None:
    capability = await add_capability(
        db_session,
        domain_name="Languages",
        canonical_name="Python",
    )
    before = await capability_count(db_session)

    result = await map_resume_skills(db_session, [skill("Python")])

    assert result.skills[0].capability_id == capability.id
    assert result.skills[0].mapping_method == "canonical_exact"
    assert result.skills[0].mapping_status == "mapped"
    assert await capability_count(db_session) == before


async def test_same_canonical_name_across_domains_stays_unmapped(db_session) -> None:
    await add_capability(
        db_session,
        domain_name="Languages",
        canonical_name="Python",
    )
    await add_capability(
        db_session,
        domain_name="Data Tools",
        canonical_name="Python",
    )
    before = await capability_count(db_session)

    result = await map_resume_skills(db_session, [skill("Python")])

    assert result.skills[0].capability_id is None
    assert result.skills[0].mapping_status == "unmapped"
    assert result.warnings == ["AMBIGUOUS_CAPABILITY_NAME:python"]
    assert await capability_count(db_session) == before


async def test_alias_exact_requires_active_alias_and_active_target(db_session) -> None:
    active = await add_capability(
        db_session,
        domain_name="Active",
        canonical_name="Python Language",
        alias="Py",
    )
    await add_capability(
        db_session,
        domain_name="Deprecated Alias",
        canonical_name="Old Python",
        alias="OldPy",
        alias_status="deprecated",
    )
    await add_capability(
        db_session,
        domain_name="Deprecated Target",
        canonical_name="Legacy Python",
        status="deprecated",
        alias="LegacyPy",
    )
    before = await capability_count(db_session)

    result = await map_resume_skills(
        db_session,
        [skill("Py"), skill("OldPy", start=10), skill("LegacyPy", start=20)],
    )
    by_name = {item.raw_name: item for item in result.skills}

    assert by_name["Py"].capability_id == active.id
    assert by_name["Py"].mapping_method == "alias_exact"
    assert by_name["OldPy"].mapping_status == "unmapped"
    assert by_name["LegacyPy"].mapping_status == "unmapped"
    assert await capability_count(db_session) == before


async def test_unmatched_skill_stays_unmapped(db_session) -> None:
    before = await capability_count(db_session)

    result = await map_resume_skills(db_session, [skill("Rust")])

    assert result.skills[0].raw_name == "Rust"
    assert result.skills[0].capability_id is None
    assert result.skills[0].mapping_method == "unmapped"
    assert await capability_count(db_session) == before


async def test_duplicate_normalized_name_prefers_strength_confidence_position(
    db_session,
) -> None:
    before = await capability_count(db_session)
    values = [
        skill("Python", strength="mention", confidence=0.99, start=0),
        skill(" python ", strength="work", confidence=0.5, start=20),
        skill("ＰＹＴＨＯＮ", strength="work", confidence=0.9, start=30),
        skill("Python", strength="work", confidence=0.9, start=10),
    ]

    result = await map_resume_skills(db_session, values)

    assert len(result.skills) == 1
    assert result.skills[0].evidence_strength == "work"
    assert result.skills[0].confidence == 0.9
    assert result.skills[0].evidence_start == 10
    assert await capability_count(db_session) == before


async def test_different_names_same_capability_keep_one_best_skill(db_session) -> None:
    capability = await add_capability(
        db_session,
        domain_name="Languages",
        canonical_name="Python",
        alias="Py",
    )
    before = await capability_count(db_session)

    result = await map_resume_skills(
        db_session,
        [
            skill("Python", strength="mention", confidence=0.99),
            skill("Py", strength="work", confidence=0.8, start=20),
        ],
    )

    assert len(result.skills) == 1
    assert result.skills[0].raw_name == "Py"
    assert result.skills[0].capability_id == capability.id
    assert result.skills[0].mapping_method == "alias_exact"
    assert await capability_count(db_session) == before


def valid_payload_with_python_evidence() -> ResumeParseResponse:
    return ResumeParseResponse.model_validate(
        {
            "schema_version": "resume_parse_v1",
            "document_language": "zh-CN",
            "summary": "具有 Python 项目经验",
            "educations": [],
            "experiences": [],
            "projects": [],
            "skills": [
                {
                    "name": "Python",
                    "proficiency": "intermediate",
                    "explicit_experience_months": 24,
                    "evidence_strength": "project",
                    "evidence_quote": "使用 Python 开发项目",
                    "confidence": 0.95,
                }
            ],
        }
    )


class FakeResponsesClient:
    def __init__(self, payload: ResumeParseResponse):
        self.payload = payload
        self.calls = 0

    async def parse_resume(self, **kwargs):
        self.calls += 1
        return LLMParseResult(
            payload=self.payload,
            response_id="resp_test",
            returned_model="fake-model",
            status="completed",
            usage={"input_tokens": 10, "output_tokens": 20, "total_tokens": 30},
            provider_attempts=1,
            response_sha256="a" * 64,
        )


@pytest.fixture
def configured_llm(monkeypatch):
    settings = SimpleNamespace(
        llm_responses_url="https://provider.test/v1/responses",
        llm_api_key=SecretStr("test-key"),
        llm_model="test-model",
    )
    monkeypatch.setattr("app.resumes.tasks.get_settings", lambda: settings)
    return settings


@pytest_asyncio.fixture
async def fake_resume_file(db_session, user, tmp_path):
    storage = FileStorage(tmp_path / "resume-task-files")
    file_id = uuid4()
    storage_key = f"resume/{file_id}.docx"
    path = storage.resolve(storage_key)
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("使用 Python 开发项目")
    document.save(path)
    content = path.read_bytes()
    stored_file = StoredFile(
        id=file_id,
        uploaded_by_user_id=user.id,
        original_name="resume.docx",
        storage_key=storage_key,
        media_type=DOCX_MEDIA_TYPE,
        extension="docx",
        size_bytes=len(content),
        sha256=sha256(content).hexdigest(),
        category="resume",
        scan_status="not_required",
        status="attached",
    )
    resume = Resume(
        id=uuid4(),
        owner_user_id=user.id,
        file_id=file_id,
        display_name="resume.docx",
        source_language="zh-CN",
        parse_status="processing",
        created_by_user_id=user.id,
    )
    db_session.add(stored_file)
    await db_session.flush()
    db_session.add(resume)
    await db_session.flush()
    return SimpleNamespace(storage=storage, stored_file=stored_file, resume=resume)


@pytest_asyncio.fixture
async def resume_run(db_session, user, fake_resume_file):
    run = ProcessingRun(
        id=uuid4(),
        run_type="parse_resume",
        subject_type="resume",
        subject_id=fake_resume_file.resume.id,
        created_by_user_id=user.id,
        owner_scope_type="user",
        owner_scope_id=user.id,
        status="pending",
        pipeline_version="resume_parse_v1",
        total_count=1,
        max_attempts=1,
        input_snapshot={
            "resume_id": str(fake_resume_file.resume.id),
            "file_id": str(fake_resume_file.stored_file.id),
        },
        result_summary={},
    )
    db_session.add(run)
    await db_session.flush()
    fake_resume_file.resume.latest_run_id = run.id
    await db_session.flush()
    return run


async def assert_failed(
    db_session,
    run,
    resume,
    *,
    code: str,
    stage: str,
    retryable: bool,
) -> None:
    await db_session.refresh(run)
    await db_session.refresh(resume)
    error = await db_session.scalar(
        select(ProcessingError).where(ProcessingError.run_id == run.id)
    )
    assert run.status == "failed"
    assert resume.parse_status == "failed"
    assert resume.latest_run_id == run.id
    assert run.error_code == code
    assert error.stage == stage
    assert error.retryable is retryable
    assert "secret" not in (run.error_message or "").lower()


async def test_parse_resume_creates_candidate_profile_and_completes_run(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    client = FakeResponsesClient(valid_payload_with_python_evidence())
    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(db_session, resume_run.id, responses_client=client)

    await db_session.refresh(resume_run)
    resume = await db_session.get(Resume, resume_run.subject_id)
    profile = await db_session.scalar(
        select(ResumeProfile).where(ResumeProfile.resume_id == resume.id)
    )
    skills = (
        await db_session.scalars(
            select(ResumeSkill).where(ResumeSkill.profile_id == profile.id)
        )
    ).all()
    assert resume.parse_status == "ready"
    assert resume.latest_run_id == resume_run.id
    assert resume_run.status == "completed"
    assert float(resume_run.progress_percent) == 100
    assert resume_run.current_stage == "completed"
    assert resume_run.result_summary["result_url"].endswith("/profiles/1")
    assert profile.status == "candidate"
    assert profile.profile_source == "extracted"
    assert profile.version_no == 1
    assert [value.raw_name for value in skills] == ["Python"]
    assert client.calls == 1


async def test_existing_extracted_profile_is_reused_without_provider_call(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    existing = ResumeProfile(
        id=uuid4(),
        resume_id=fake_resume_file.resume.id,
        version_no=1,
        extraction_version="resume_parse_v1",
        profile_source="extracted",
        extracted_text="使用 Python 开发项目",
        text_extraction_method="docx",
        structured_payload={},
        status="candidate",
        created_by_run_id=resume_run.id,
        created_by_user_id=resume_run.created_by_user_id,
    )
    db_session.add(existing)
    await db_session.flush()
    resume_run.status = "failed"
    retry = ProcessingRun(
        id=uuid4(),
        run_type="parse_resume",
        subject_type="resume",
        subject_id=fake_resume_file.resume.id,
        retry_of_run_id=resume_run.id,
        created_by_user_id=resume_run.created_by_user_id,
        owner_scope_type="user",
        owner_scope_id=resume_run.owner_scope_id,
        status="pending",
        pipeline_version="resume_parse_v1",
        total_count=1,
        max_attempts=1,
        input_snapshot=dict(resume_run.input_snapshot),
        result_summary={},
    )
    db_session.add(retry)
    await db_session.flush()
    fake_resume_file.resume.latest_run_id = retry.id
    await db_session.flush()
    client = FakeResponsesClient(valid_payload_with_python_evidence())
    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(db_session, retry.id, responses_client=client)

    await db_session.refresh(retry)
    assert retry.status == "completed"
    assert retry.result_summary["profile_id"] == str(existing.id)
    assert await db_session.scalar(select(func.count()).select_from(ResumeProfile)) == 1
    assert client.calls == 0


async def test_provider_call_runs_without_active_transaction(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    states = []

    class TransactionCheckingClient(FakeResponsesClient):
        async def parse_resume(self, **kwargs):
            states.append(db_session.in_transaction())
            assert resume_run.status == "running"
            assert resume_run.current_stage == "call_llm"
            return await super().parse_resume(**kwargs)

    client = TransactionCheckingClient(valid_payload_with_python_evidence())
    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(db_session, resume_run.id, responses_client=client)

    assert states == [False]


async def test_pending_cancelled_run_does_not_parse(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    client = FakeResponsesClient(valid_payload_with_python_evidence())
    resume_run.status = "cancelled"
    resume_run.cancel_requested = True
    resume_run.completed_at = datetime.now(UTC)
    await db_session.flush()
    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(db_session, resume_run.id, responses_client=client)

    await db_session.refresh(fake_resume_file.resume)
    assert fake_resume_file.resume.parse_status == "uploaded"
    assert client.calls == 0


async def test_cancel_requested_after_provider_discards_result(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    class CancellingClient(FakeResponsesClient):
        async def parse_resume(self, **kwargs):
            result = await super().parse_resume(**kwargs)
            resume_run.cancel_requested = True
            resume_run.status = "cancel_requested"
            await db_session.commit()
            return result

    client = CancellingClient(valid_payload_with_python_evidence())
    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(db_session, resume_run.id, responses_client=client)

    await db_session.refresh(resume_run)
    await db_session.refresh(fake_resume_file.resume)
    assert resume_run.status == "cancelled"
    assert fake_resume_file.resume.parse_status == "uploaded"
    assert await db_session.scalar(select(func.count()).select_from(ResumeProfile)) == 0


async def test_document_error_marks_run_and_resume_failed(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    document = Document()
    document.save(fake_resume_file.storage.resolve(fake_resume_file.stored_file.storage_key))
    client = FakeResponsesClient(valid_payload_with_python_evidence())
    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(db_session, resume_run.id, responses_client=client)

    await assert_failed(
        db_session,
        resume_run,
        fake_resume_file.resume,
        code="RESUME_TEXT_EMPTY",
        stage="extract_text",
        retryable=False,
    )
    assert client.calls == 0


async def test_llm_error_marks_safe_processing_error(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    class FailingClient:
        async def parse_resume(self, **kwargs):
            raise ResumeLLMError("LLM_RATE_LIMITED", "request", True, 429)

    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(
        db_session,
        resume_run.id,
        responses_client=FailingClient(),
    )

    await assert_failed(
        db_session,
        resume_run,
        fake_resume_file.resume,
        code="LLM_RATE_LIMITED",
        stage="call_llm",
        retryable=True,
    )


async def test_unconfigured_llm_marks_safe_processing_error(
    db_session,
    resume_run,
    fake_resume_file,
    monkeypatch,
) -> None:
    monkeypatch.setattr(
        "app.resumes.tasks.get_settings",
        lambda: SimpleNamespace(
            llm_responses_url=None,
            llm_api_key=None,
            llm_model=None,
        ),
    )
    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(
        db_session,
        resume_run.id,
        responses_client=FakeResponsesClient(valid_payload_with_python_evidence()),
    )

    await assert_failed(
        db_session,
        resume_run,
        fake_resume_file.resume,
        code="LLM_NOT_CONFIGURED",
        stage="call_llm",
        retryable=False,
    )


async def test_all_invalid_evidence_marks_failed(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    values = valid_payload_with_python_evidence().model_dump()
    values["skills"][0]["evidence_quote"] = "不存在的 Java 证据"
    client = FakeResponsesClient(ResumeParseResponse.model_validate(values))
    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)

    await run_parse_resume(db_session, resume_run.id, responses_client=client)

    await assert_failed(
        db_session,
        resume_run,
        fake_resume_file.resume,
        code="RESUME_EVIDENCE_EMPTY",
        stage="validate_evidence",
        retryable=False,
    )


async def test_persistence_failure_rolls_back_profile_and_marks_failed(
    db_session,
    resume_run,
    fake_resume_file,
    configured_llm,
    monkeypatch,
) -> None:
    async def fail_persistence(*args, **kwargs):
        raise SQLAlchemyError("secret database detail")

    monkeypatch.setattr("app.resumes.tasks.storage", fake_resume_file.storage)
    monkeypatch.setattr(
        "app.resumes.tasks.persist_extracted_profile",
        fail_persistence,
    )

    await run_parse_resume(
        db_session,
        resume_run.id,
        responses_client=FakeResponsesClient(valid_payload_with_python_evidence()),
    )

    await assert_failed(
        db_session,
        resume_run,
        fake_resume_file.resume,
        code="RESUME_PERSISTENCE_FAILED",
        stage="persist_profile",
        retryable=True,
    )
    assert await db_session.scalar(select(func.count()).select_from(ResumeProfile)) == 0
    assert await db_session.scalar(select(func.count()).select_from(ResumeSkill)) == 0


async def test_retry_run_uses_same_generic_task_name(
    db_session,
    resume_run,
    user,
    monkeypatch,
) -> None:
    sent = []
    resume_run.status = "failed"
    await db_session.flush()

    def send_task(name, args):
        sent.append((name, args))
        return SimpleNamespace(id="retry-task-id")

    monkeypatch.setattr("app.processing.service.celery_app.send_task", send_task)

    new_run = await retry_run(db_session, resume_run, user)

    assert sent == [("app.parse_resume", [str(new_run.id)])]
